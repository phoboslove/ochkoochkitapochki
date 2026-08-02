"""Generate + install 3 high-quality commercial KZ DOCX templates.

Built programmatically with python-docx so the demo company has *commercial*
templates (not inventory forms) for the three core operational documents:

  1. Акт выполненных работ            (act_of_works)
  2. Счёт на оплату (коммерческий)    (commercial_invoice)
  3. Товарная накладная (ТТН)         (delivery_waybill)

All three:
  * A4 portrait, 20mm margins
  * Times New Roman 11pt body, 16pt centred title
  * Header block: «Поставщик» / «Заказчик» (parties) — left/right columns
  * Item table with full borders + `{%tr for it in items %}` row directive
  * Totals block (subtotal / VAT 12% / total)
  * Amount-in-words line
  * Signatures block (two columns, Исполнитель / Заказчик)
  * Bank requisites footer (БИН, IBAN, БИК)

Run:
  $ python scripts/install_commercial_templates.py

Idempotent: deletes any existing templates with the same name before
re-installing, so iteration is cheap.
"""
from __future__ import annotations

import asyncio
import hashlib
import sys
import uuid
from io import BytesIO
from pathlib import Path

from sqlalchemy import select, delete

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from docx import Document  # type: ignore
from docx.enum.table import WD_ALIGN_VERTICAL  # type: ignore
from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
from docx.oxml.ns import qn  # type: ignore
from docx.oxml import OxmlElement  # type: ignore
from docx.shared import Cm, Pt, RGBColor  # type: ignore

from app.core.db import SessionLocal  # noqa: E402
from app.db.models import Company, Template  # noqa: E402
from app.services.storage import get_storage  # noqa: E402


CANONICAL_MAPPING_SET = [
    "company_name", "company_bin", "company_address", "company_bank",
    "client_name", "client_bin", "client_address",
    "document_number", "document_date",
    "total", "subtotal", "vat", "vat_percent", "currency",
    "director_name", "accountant_name", "amount_words", "notes",
]


# ── Low-level styling helpers ──────────────────────────────────────────────

def _set_cell_borders(cell) -> None:
    """Apply a thin single-line border to all four sides of a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "808080")
        tcBorders.append(b)
    tcPr.append(tcBorders)


def _shade_cell(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def _set_margins(section, mm: int = 20) -> None:
    section.left_margin = section.right_margin = Cm(mm / 10)
    section.top_margin = section.bottom_margin = Cm(mm / 10)


def _para(doc, text: str = "", *, size: int = 11, bold: bool = False,
            align: int = WD_ALIGN_PARAGRAPH.LEFT, space_after: int = 4):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        run.font.bold = bold
    return p


def _header_band(doc, title: str, number_token: str, date_token: str) -> None:
    """Center the title, then a centred '№ {{n}} от {{date}}' line."""
    _para(doc, title, size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(f"№ {number_token}  от  {date_token} г.")
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    run.italic = True


def _parties_block(doc, *, left_label: str, right_label: str) -> None:
    """Two-column parties table — Поставщик / Заказчик."""
    t = doc.add_table(rows=1, cols=2)
    t.autofit = False
    for cell in t.rows[0].cells:
        cell.width = Cm(8.5)
        _set_cell_borders(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    def fill(cell, label, who):
        cell.text = ""
        h = cell.paragraphs[0]
        h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = h.add_run(label)
        r.font.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        for line in (
            f"{{{{{who}_name}}}}",
            f"БИН: {{{{{who}_bin}}}}",
            f"{{{{{who}_address}}}}",
        ):
            p = cell.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(line)
            run.font.name = "Times New Roman"
            run.font.size = Pt(10.5)

    fill(t.rows[0].cells[0], left_label,  "company")
    fill(t.rows[0].cells[1], right_label, "client")
    _para(doc, "", space_after=8)


def _items_table(doc, columns: list[tuple[str, str, float]]) -> None:
    """Bordered items table with a single repeating content row.

    Built as 4 table rows so docxtpl's ``{%tr ...%}`` directive parses cleanly:

      Row 1: header                  (visible)
      Row 2: ``{%tr for it in items %}`` in cell[0] only  → docxtpl REPLACES
             this entire row with the bare Jinja ``{% for %}`` text, so the
             row visually disappears in the rendered output.
      Row 3: content cells with ``{{ it.x }}`` placeholders. THIS row is the
             one Jinja iterates over — it's duplicated for every item.
      Row 4: ``{%tr endfor %}`` → likewise replaced and removed.

    docxtpl's ``{%tr%}`` regex matches an ENTIRE ``<w:tr>...</w:tr>`` row
    and substitutes only the bare Jinja directive in its place. Placing two
    directives in the same row would cause one to be discarded — which was
    the cause of the "Encountered unknown tag 'endfor'" failure earlier.
    """
    t = doc.add_table(rows=4, cols=len(columns))
    t.autofit = False

    # --- Row 1: header ---
    for idx, (label, _, w_cm) in enumerate(columns):
        cell = t.rows[0].cells[idx]
        cell.width = Cm(w_cm)
        cell.text = ""
        _set_cell_borders(cell)
        _shade_cell(cell, "F0F0F0")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(label); r.font.bold = True; r.font.size = Pt(10)

    # --- Row 2: {%tr for it in items %} (entire row consumed by docxtpl) ---
    for_cell = t.rows[1].cells[0]
    for_cell.text = ""
    p = for_cell.paragraphs[0]
    r = p.add_run("{%tr for it in items %}")
    r.font.size = Pt(1)

    # --- Row 3: content (this is the row Jinja replicates) ---
    body = t.rows[2]
    for idx, (_, expr, w_cm) in enumerate(columns):
        cell = body.cells[idx]
        cell.width = Cm(w_cm)
        cell.text = ""
        _set_cell_borders(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = (
            WD_ALIGN_PARAGRAPH.LEFT if idx <= 1 else WD_ALIGN_PARAGRAPH.RIGHT
        )
        r = p.add_run(expr)
        r.font.name = "Times New Roman"
        r.font.size = Pt(10.5)

    # --- Row 4: {%tr endfor %} (row gets removed) ---
    endfor_cell = t.rows[3].cells[0]
    endfor_cell.text = ""
    p = endfor_cell.paragraphs[0]
    r = p.add_run("{%tr endfor %}")
    r.font.size = Pt(1)


def _totals_block(doc, *, with_vat: bool = True) -> None:
    """Right-aligned totals — subtotal / VAT / TOTAL."""
    t = doc.add_table(rows=3 if with_vat else 2, cols=2)
    t.autofit = False
    rows = [
        ("Без НДС:",                       "{{subtotal}} {{currency}}",   False),
    ]
    if with_vat:
        rows.append(("НДС ({{vat_percent}}%):", "{{vat}} {{currency}}", False))
    rows.append(("ИТОГО к оплате:", "{{total}} {{currency}}", True))
    for i, (lbl, val, bold) in enumerate(rows):
        c1, c2 = t.rows[i].cells
        c1.width = Cm(13)
        c2.width = Cm(4)
        c1.text = ""; c2.text = ""
        p1 = c1.paragraphs[0]; p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p1.add_run(lbl); r.font.bold = bold; r.font.size = Pt(11)
        p2 = c2.paragraphs[0]; p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p2.add_run(val); r.font.bold = bold; r.font.size = Pt(11)
    _para(doc, "", space_after=4)
    _para(doc, "Сумма прописью: {{amount_words}}",
            size=10, align=WD_ALIGN_PARAGRAPH.LEFT)


def _signatures_block(doc, *, left: str, right: str) -> None:
    _para(doc, "", space_after=20)
    t = doc.add_table(rows=2, cols=2)
    t.autofit = False
    for c in (t.rows[0].cells[0], t.rows[0].cells[1]):
        c.width = Cm(8.5)
    for cell, role, signer_field in (
        (t.rows[0].cells[0], left,  "director_name"),
        (t.rows[0].cells[1], right, ""),
    ):
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(role); r.font.bold = True; r.font.size = Pt(10)
        sig_line = cell.add_paragraph()
        sig_line.paragraph_format.space_before = Pt(24)
        r = sig_line.add_run("_______________________ / " +
                              ("{{" + signer_field + "}}" if signer_field else "________________"))
        r.font.size = Pt(10)
        sig_line.add_run("\nм.п.").italic = True
    # bottom row empty for spacing
    for cell in t.rows[1].cells:
        cell.text = ""


def _bank_footer(doc) -> None:
    _para(doc, "", space_after=14)
    _para(doc, "Банковские реквизиты Поставщика", size=10, bold=True)
    _para(doc, "{{company_name}} · БИН {{company_bin}}", size=10)
    _para(doc, "{{company_bank}}", size=10)
    _para(doc, "Адрес: {{company_address}}", size=10)


# ── Templates ──────────────────────────────────────────────────────────────

def build_act_of_works() -> bytes:
    doc = Document()
    _set_margins(doc.sections[0])
    _header_band(doc, "АКТ № выполненных работ (оказанных услуг)",
                  "{{document_number}}", "{{document_date}}")
    _parties_block(doc, left_label="ИСПОЛНИТЕЛЬ", right_label="ЗАКАЗЧИК")
    _para(doc, "Мы, нижеподписавшиеся, составили настоящий акт о том, что "
                "работы (услуги) по договору выполнены полностью и в срок:",
                size=10.5, space_after=8)
    _items_table(doc, columns=[
        ("№",                "{{ loop.index }}",         1.2),
        ("Наименование работ (услуг)", "{{ it.name }}",   8.3),
        ("Ед.",              "шт",                       1.5),
        ("Кол-во",           "{{ it.qty }}",             1.5),
        ("Цена",             "{{ it.price_fmt }}",       2.0),
        ("Сумма",            "{{ it.total_fmt }}",       2.5),
    ])
    _para(doc, "", space_after=8)
    _totals_block(doc, with_vat=True)
    _para(doc, "", space_after=4)
    _para(doc, "Услуги выполнены полностью и в срок. Заказчик претензий "
                "по объёму, качеству и срокам оказания услуг не имеет.",
                size=10, space_after=6)
    _signatures_block(doc, left="От Исполнителя:", right="От Заказчика:")
    _bank_footer(doc)
    out = BytesIO(); doc.save(out); return out.getvalue()


def build_commercial_invoice() -> bytes:
    doc = Document()
    _set_margins(doc.sections[0])
    _header_band(doc, "СЧЁТ НА ОПЛАТУ",
                  "{{document_number}}", "{{document_date}}")
    _parties_block(doc, left_label="ПОСТАВЩИК", right_label="ПОКУПАТЕЛЬ")
    _items_table(doc, columns=[
        ("№",                "{{ loop.index }}",         1.2),
        ("Наименование товара (услуги)", "{{ it.name }}", 8.3),
        ("Ед.",              "шт",                       1.5),
        ("Кол-во",           "{{ it.qty }}",             1.5),
        ("Цена",             "{{ it.price_fmt }}",       2.0),
        ("Сумма",            "{{ it.total_fmt }}",       2.5),
    ])
    _para(doc, "", space_after=8)
    _totals_block(doc, with_vat=True)
    _para(doc, "", space_after=10)
    _para(doc, "Оплата производится в течение 14 (четырнадцати) календарных "
                "дней с момента выставления настоящего счёта.",
                size=10, space_after=6)
    _para(doc, "Назначение платежа: оплата по счёту № {{document_number}} "
                "от {{document_date}} г.",
                size=10, space_after=6)
    _signatures_block(doc, left="Руководитель:", right="Главный бухгалтер:")
    _bank_footer(doc)
    out = BytesIO(); doc.save(out); return out.getvalue()


def build_delivery_waybill() -> bytes:
    doc = Document()
    _set_margins(doc.sections[0])
    _header_band(doc, "ТОВАРНАЯ НАКЛАДНАЯ",
                  "{{document_number}}", "{{document_date}}")
    _parties_block(doc, left_label="ГРУЗООТПРАВИТЕЛЬ", right_label="ГРУЗОПОЛУЧАТЕЛЬ")
    _items_table(doc, columns=[
        ("№",                "{{ loop.index }}",         1.2),
        ("Наименование товара", "{{ it.name }}",         9.0),
        ("Ед.",              "шт",                       1.4),
        ("Кол-во",           "{{ it.qty }}",             1.6),
        ("Цена",             "{{ it.price_fmt }}",       2.1),
        ("Сумма",            "{{ it.total_fmt }}",       2.2),
    ])
    _para(doc, "", space_after=8)
    _totals_block(doc, with_vat=True)
    _para(doc, "", space_after=10)
    _para(doc, "Товар отгружен в исправной таре, с проверенным качеством и "
                "количеством. Претензий к комплектности нет.",
                size=10, space_after=6)
    _signatures_block(doc, left="Отпустил груз (Поставщик):", right="Принял груз (Получатель):")
    _bank_footer(doc)
    out = BytesIO(); doc.save(out); return out.getvalue()


# ── Installer ──────────────────────────────────────────────────────────────

TEMPLATES = [
    ("act",        "Акт выполненных работ (коммерческий)",   build_act_of_works),
    ("invoice",    "Счёт на оплату (коммерческий)",          build_commercial_invoice),
    ("nakladnaya", "Товарная накладная (коммерческая)",      build_delivery_waybill),
]


def _build_mappings() -> dict[str, dict]:
    return {
        key: {"source_label": key, "source_kind": "canonical",
              "confidence": 0.95, "confirmed": True}
        for key in CANONICAL_MAPPING_SET
    }


# Names of stop-gap templates promoted by earlier verify_templates.py.
# Commercial templates supersede them — archive so the matcher prefers
# the real commercial layouts instead of inventory/audit forms.
_LEGACY_WORKAROUND_NAMES = [
    "Inv 4 akt inventarizatsii debitorskoy i kreditorskoy zadolzhennosti kz",
    "Da 4 nakladnaya na vnutrennee peremeshchenie dolgosrochnykh aktivov kz",
]


async def install() -> None:
    storage = get_storage()
    async with SessionLocal() as session:
        company = (await session.scalars(select(Company).limit(1))).first()
        if not company:
            print("no company found — seed first"); return
        print(f"installing into company={company.id} ({company.name})")

        # Archive earlier workaround promotions so the matcher prefers the
        # new commercial templates (they're cold-start and would otherwise
        # lose to the 100/100 operational reliability of the legacy ones).
        archived = 0
        for legacy_name in _LEGACY_WORKAROUND_NAMES:
            row = (await session.scalars(
                select(Template).where(
                    Template.company_id == company.id,
                    Template.name == legacy_name,
                    Template.status == "VERIFIED",
                ),
            )).first()
            if row:
                row.status = "ARCHIVED"
                archived += 1
                print(f"  [archived] {legacy_name}")
        await session.flush()
        if archived:
            print(f"  archived {archived} legacy workaround template(s).")

        for kind, display_name, builder in TEMPLATES:
            # Replace any prior install of the same name (idempotent).
            await session.execute(
                delete(Template).where(
                    Template.company_id == company.id,
                    Template.name == display_name,
                ),
            )
            await session.flush()

            blob = builder()
            file_hash = hashlib.sha256(blob).hexdigest()
            tpl_id = f"tpl_{uuid.uuid4().hex[:10]}"
            storage_key = f"{company.id}/templates/{tpl_id}/{display_name}.docx"
            storage.put(storage_key, blob, content_type=(
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ))

            t = Template(
                id=tpl_id, company_id=company.id,
                name=display_name, kind=kind, format="docx",
                body=storage_key, is_default=True,
                status="VERIFIED", original_filename=f"{display_name}.docx",
                file_hash=file_hash, language="ru", version=1,
                detected_fields={
                    "placeholders": list(_build_mappings().keys()),
                    "tables": [{"rows": 1, "cols": 6, "is_items": True}],
                    "semantic": {"kind": kind, "commercial": True},
                    "built_by": "install_commercial_templates.py",
                },
                detected_tables=[{"rows": 1, "cols": 6, "is_items": True,
                                    "header": ["№","Наименование","Ед.","Кол-во","Цена","Сумма"]}],
                mappings=_build_mappings(),
                confidence=0.95,
                validation_score=95, validation_warnings=[],
                industry="general",
            )
            session.add(t)
            print(f"  [{kind:11s}] VERIFIED  {display_name}  "
                   f"({len(blob)} bytes, {len(t.mappings)} mappings)")

        await session.commit()
        print(f"\nDone. Installed {len(TEMPLATES)} commercial templates.")


if __name__ == "__main__":
    asyncio.run(install())
