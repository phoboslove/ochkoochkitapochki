"""Generate + install 6 VERIFIED KZ legal/HR templates.

Companion to ``install_commercial_templates.py`` (same conventions: raw
python-docx authoring with literal ``{{placeholder}}`` / ``{%tr %}`` runs,
rendered later by docxtpl). Kept as a separate script/kind-set rather than
extending the commercial one because these six documents needed brand-new
``kind`` values and canonical context fields (see
``app/services/documents/generation/context_builder.py``) that the three
commercial templates never touched.

  1. Доверенность на получение ТМЦ        (trust_letter)
  2. Договор оказания услуг                (contract_services)
  3. Договор поставки                      (contract_supply)
  4. Акт сверки взаиморасчётов             (act_reconciliation)
  5. Приказ о приёме на работу             (hr_order)
  6. Трудовой договор                      (employment_contract)

Wording is original — structured off the mandatory-requisite conventions in
adilet.zan.kz's published forms and general RK contracting/labour practice,
not copied from any commercial site.

Run:
  $ python scripts/install_kz_legal_templates.py

Idempotent: deletes any existing templates with the same name before
re-installing, so iteration is cheap.
"""
from __future__ import annotations

import asyncio
import hashlib
import sys
import uuid
from io import BytesIO
from pathlib import Path

from sqlalchemy import select, delete

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from docx import Document  # type: ignore
from docx.enum.table import WD_ALIGN_VERTICAL  # type: ignore
from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
from docx.oxml.ns import qn  # type: ignore
from docx.oxml import OxmlElement  # type: ignore
from docx.shared import Cm, Pt, RGBColor  # type: ignore

from app.core.db import SessionLocal  # noqa: E402
from app.db.models import Company, Template  # noqa: E402
from app.services.storage import get_storage  # noqa: E402


# Canonical mapping set per kind — used to pre-fill Template.mappings with
# confirmed=True, same pattern as install_commercial_templates.py's single
# shared CANONICAL_MAPPING_SET, but split per-kind here since the six
# documents don't share one field vocabulary the way act/invoice/nakladnaya do.
_BASE_FIELDS = [
    "company_name", "company_bin", "company_address", "company_bank",
    "document_number", "document_date", "director_name", "accountant_name",
]

CANONICAL_MAPPING_SETS: dict[str, list[str]] = {
    "trust_letter": _BASE_FIELDS + [
        "employee_name", "employee_position", "employee_iin", "employee_id_doc",
        "valid_until", "from_name", "from_bin", "items",
    ],
    "contract_services": _BASE_FIELDS + [
        "client_name", "client_bin", "client_address", "client_director_name",
        "service_description", "total", "currency", "payment_terms",
        "start_date", "end_date", "city",
    ],
    "contract_supply": _BASE_FIELDS + [
        "client_name", "client_bin", "client_address", "client_director_name",
        "items", "total", "subtotal", "vat", "vat_percent", "currency",
        "payment_terms", "delivery_terms", "delivery_address", "delivery_date", "city",
    ],
    "act_reconciliation": _BASE_FIELDS + [
        "client_name", "client_bin", "client_address", "client_director_name",
        "period_start", "period_end", "opening_balance", "closing_balance",
        "closing_balance_words", "operations", "currency",
    ],
    "hr_order": _BASE_FIELDS + [
        "order_number", "employee_name", "employee_iin", "employee_position",
        "employee_department", "hire_date", "salary", "probation_period", "city",
    ],
    "employment_contract": _BASE_FIELDS + [
        "employee_name", "employee_iin", "employee_address", "employee_id_doc",
        "employee_position", "employee_department", "work_start_date",
        "contract_term", "salary", "pay_dates", "work_schedule",
        "probation_period", "vacation_days", "city",
    ],
}


# ── Low-level styling helpers (same conventions as install_commercial_templates.py) ──

def _set_cell_borders(cell) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "808080")
        tcBorders.append(b)
    tcPr.append(tcBorders)


def _shade_cell(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def _set_margins(section, mm: int = 20) -> None:
    section.left_margin = section.right_margin = Cm(mm / 10)
    section.top_margin = section.bottom_margin = Cm(mm / 10)


def _para(doc, text: str = "", *, size: float = 11, bold: bool = False,
          italic: bool = False, align: int = WD_ALIGN_PARAGRAPH.LEFT,
          space_after: int = 4, space_before: int = 0):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if text:
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
    return p


def _title_block(doc, title: str, number_token: str, date_token: str,
                  subtitle: str | None = None) -> None:
    _para(doc, title, size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    if subtitle:
        _para(doc, subtitle, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run(f"№ {number_token}  от  {date_token} г.")
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    run.italic = True


def _section_heading(doc, text: str) -> None:
    _para(doc, text, size=11.5, bold=True, space_before=10, space_after=4)


def _parties_block(doc, *, left_label: str, right_label: str,
                    left_who: str = "company", right_who: str = "client",
                    left_extra: list[str] | None = None,
                    right_extra: list[str] | None = None,
                    left_base: list[str] | None = None,
                    right_base: list[str] | None = None) -> None:
    """Two-column requisites table — generalizes install_commercial_templates.py's
    version to accept arbitrary placeholder prefixes and extra lines (needed
    for e.g. employee/employer blocks that don't fit the company/client shape).

    ``left_base``/``right_base`` override the default name/БИН/address triplet
    entirely — needed for an "employee" column, which has no БИН.
    """
    t = doc.add_table(rows=1, cols=2)
    t.autofit = False
    for cell in t.rows[0].cells:
        cell.width = Cm(8.5)
        _set_cell_borders(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    def fill(cell, label, who, extra, base):
        cell.text = ""
        h = cell.paragraphs[0]
        h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = h.add_run(label)
        r.font.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        lines = list(base) if base is not None else [
            f"{{{{{who}_name}}}}", f"БИН: {{{{{who}_bin}}}}", f"{{{{{who}_address}}}}",
        ]
        lines.extend(extra or [])
        for line in lines:
            p = cell.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(line)
            run.font.name = "Times New Roman"
            run.font.size = Pt(10.5)

    fill(t.rows[0].cells[0], left_label, left_who, left_extra, left_base)
    fill(t.rows[0].cells[1], right_label, right_who, right_extra, right_base)
    _para(doc, "", space_after=8)


def _table(doc, *, loop_var: str, columns: list[tuple[str, str, float]]) -> None:
    """Bordered table with a single repeating content row, parametrized by
    the Jinja loop variable name (``items`` for goods/services rows,
    ``operations`` for the reconciliation-act ledger). See
    install_commercial_templates.py's ``_items_table`` docstring for why
    the ``{%tr %}``/``{%tr endfor %}`` directives each need their own row.
    """
    t = doc.add_table(rows=4, cols=len(columns))
    t.autofit = False

    for idx, (label, _, w_cm) in enumerate(columns):
        cell = t.rows[0].cells[idx]
        cell.width = Cm(w_cm)
        cell.text = ""
        _set_cell_borders(cell)
        _shade_cell(cell, "F0F0F0")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(label); r.font.bold = True; r.font.size = Pt(10)

    for_cell = t.rows[1].cells[0]
    for_cell.text = ""
    p = for_cell.paragraphs[0]
    r = p.add_run(f"{{%tr for it in {loop_var} %}}")
    r.font.size = Pt(1)

    body = t.rows[2]
    for idx, (_, expr, w_cm) in enumerate(columns):
        cell = body.cells[idx]
        cell.width = Cm(w_cm)
        cell.text = ""
        _set_cell_borders(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx <= 1 else WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run(expr)
        r.font.name = "Times New Roman"
        r.font.size = Pt(10.5)

    endfor_cell = t.rows[3].cells[0]
    endfor_cell.text = ""
    p = endfor_cell.paragraphs[0]
    r = p.add_run("{%tr endfor %}")
    r.font.size = Pt(1)


def _totals_block(doc, *, with_vat: bool = True) -> None:
    t = doc.add_table(rows=3 if with_vat else 2, cols=2)
    t.autofit = False
    rows = [("Без НДС:", "{{subtotal}} {{currency}}", False)]
    if with_vat:
        rows.append(("НДС ({{vat_percent}}%):", "{{vat}} {{currency}}", False))
    rows.append(("ИТОГО:", "{{total}} {{currency}}", True))
    for i, (lbl, val, bold) in enumerate(rows):
        c1, c2 = t.rows[i].cells
        c1.width = Cm(13); c2.width = Cm(4)
        c1.text = ""; c2.text = ""
        p1 = c1.paragraphs[0]; p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p1.add_run(lbl); r.font.bold = bold; r.font.size = Pt(11)
        p2 = c2.paragraphs[0]; p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p2.add_run(val); r.font.bold = bold; r.font.size = Pt(11)
    _para(doc, "", space_after=4)


def _signatures_block(doc, pairs: list[tuple[str, str]]) -> None:
    """``pairs`` = [(role_label, signer_placeholder_or_blank), ...] — 2 or 3 columns.

    Single-row table — the signature line's own space_before already gives
    the visual gap under the role label, so a second, permanently-blank row
    isn't needed for spacing. (It used to be here purely for whitespace and
    had no other purpose, but quality.py's generic post-render table scanner
    can't tell "genuinely blank by design" apart from "leftover unconsumed
    {%tr%} repeating-row" — a real signal on tables that DO use {%tr%} — so
    it flagged this static row as "Table N contains an empty row" on every
    hr_order/employment_contract despite there being no data-loss.)
    """
    _para(doc, "", space_after=16)
    t = doc.add_table(rows=1, cols=len(pairs))
    t.autofit = False
    width = Cm(17 / len(pairs))
    for idx, (role, signer_field) in enumerate(pairs):
        cell = t.rows[0].cells[idx]
        cell.width = width
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(role); r.font.bold = True; r.font.size = Pt(10)
        sig_line = cell.add_paragraph()
        sig_line.paragraph_format.space_before = Pt(24)
        r = sig_line.add_run("_______________________ / " +
                              ("{{" + signer_field + "}}" if signer_field else "________________"))
        r.font.size = Pt(10)
        sig_line.add_run("\nм.п.").italic = True


# ── Templates ──────────────────────────────────────────────────────────────

def build_trust_letter() -> bytes:
    doc = Document()
    _set_margins(doc.sections[0])
    _title_block(doc, "ДОВЕРЕННОСТЬ", "{{document_number}}", "{{document_date}}")
    _para(doc,
          "{{company_name}} (БИН {{company_bin}}), в лице {{director_name}}, "
          "действующего на основании Устава, настоящей доверенностью уполномочивает:",
          size=10.5, space_after=8)
    _para(doc, "{{employee_name}}", size=12, bold=True, space_after=2)
    _para(doc, "Должность: {{employee_position}}", size=10.5, space_after=2)
    _para(doc, "ИИН: {{employee_iin}}  ·  Документ, удостоверяющий личность: {{employee_id_doc}}",
          size=10.5, space_after=10)
    _para(doc,
          "на получение товарно-материальных ценностей от {{from_name}} "
          "(БИН {{from_bin}}) по следующей номенклатуре:",
          size=10.5, space_after=8)
    _table(doc, loop_var="items", columns=[
        ("№",                "{{ loop.index }}", 1.5),
        ("Наименование",     "{{ it.name }}",     10.0),
        ("Ед.",              "шт",                2.0),
        ("Кол-во",           "{{ it.qty }}",       3.5),
    ])
    _para(doc, "", space_after=8)
    _para(doc, "Доверенность действительна до {{valid_until}} г.", size=10.5, space_after=12)
    _para(doc, "Подпись доверенного лица удостоверяю:", size=10, italic=True, space_after=4)
    _signatures_block(doc, [
        ("Руководитель:", "director_name"),
        ("Главный бухгалтер:", "accountant_name"),
    ])
    out = BytesIO(); doc.save(out); return out.getvalue()


_SERVICE_CONTRACT_BOILERPLATE = [
    ("4. ПРАВА И ОБЯЗАННОСТИ СТОРОН",
     "4.1. Исполнитель обязуется оказать услуги надлежащего качества и в "
     "согласованный срок. 4.2. Заказчик обязуется предоставить Исполнителю "
     "информацию и материалы, необходимые для оказания услуг, и своевременно "
     "оплатить оказанные услуги. 4.3. Заказчик вправе запрашивать информацию "
     "о ходе оказания услуг."),
    ("5. ОТВЕТСТВЕННОСТЬ СТОРОН",
     "5.1. За неисполнение или ненадлежащее исполнение обязательств по "
     "настоящему Договору Стороны несут ответственность в соответствии с "
     "законодательством Республики Казахстан. 5.2. Сторона, не исполнившая "
     "обязательство вследствие непреодолимой силы, освобождается от "
     "ответственности на период действия таких обстоятельств."),
    ("6. ЗАКЛЮЧИТЕЛЬНЫЕ ПОЛОЖЕНИЯ",
     "6.1. Все изменения и дополнения к настоящему Договору действительны "
     "при условии составления в письменной форме и подписания обеими "
     "Сторонами. 6.2. Споры, возникающие при исполнении настоящего Договора, "
     "разрешаются путём переговоров, а при недостижении согласия — в "
     "судебном порядке в соответствии с законодательством Республики "
     "Казахстан. 6.3. Договор составлен в двух экземплярах, имеющих "
     "одинаковую юридическую силу, по одному для каждой из Сторон."),
]

_SUPPLY_CONTRACT_BOILERPLATE = [
    ("4. ПРАВА И ОБЯЗАННОСТИ СТОРОН",
     "4.1. Поставщик обязуется передать товар надлежащего качества, "
     "комплектности и в согласованный срок. 4.2. Покупатель обязуется "
     "принять товар и произвести оплату в порядке, установленном настоящим "
     "Договором. 4.3. Покупатель вправе отказаться от приёмки товара, не "
     "соответствующего условиям Договора, с составлением акта."),
    ("5. ОТВЕТСТВЕННОСТЬ СТОРОН",
     "5.1. За неисполнение или ненадлежащее исполнение обязательств по "
     "настоящему Договору Стороны несут ответственность в соответствии с "
     "законодательством Республики Казахстан. 5.2. Сторона, не исполнившая "
     "обязательство вследствие непреодолимой силы, освобождается от "
     "ответственности на период действия таких обстоятельств."),
    ("6. ЗАКЛЮЧИТЕЛЬНЫЕ ПОЛОЖЕНИЯ",
     "6.1. Все изменения и дополнения к настоящему Договору действительны "
     "при условии составления в письменной форме и подписания обеими "
     "Сторонами. 6.2. Споры, возникающие при исполнении настоящего Договора, "
     "разрешаются путём переговоров, а при недостижении согласия — в "
     "судебном порядке в соответствии с законодательством Республики "
     "Казахстан. 6.3. Договор составлен в двух экземплярах, имеющих "
     "одинаковую юридическую силу, по одному для каждой из Сторон."),
]


def build_service_contract() -> bytes:
    doc = Document()
    _set_margins(doc.sections[0])
    _title_block(doc, "ДОГОВОР ОКАЗАНИЯ УСЛУГ", "{{document_number}}", "{{document_date}}")
    _para(doc, "г. {{city}}", size=10.5, space_after=10)
    _para(doc,
          "{{company_name}}, именуемое в дальнейшем «Исполнитель», в лице "
          "{{director_name}}, действующего на основании Устава, с одной "
          "стороны, и {{client_name}}, именуемое в дальнейшем «Заказчик», в "
          "лице {{client_director_name}}, действующего на основании Устава, "
          "с другой стороны, совместно именуемые «Стороны», заключили "
          "настоящий Договор о нижеследующем:",
          size=10.5, space_after=10)

    _section_heading(doc, "1. ПРЕДМЕТ ДОГОВОРА")
    _para(doc,
          "1.1. Исполнитель обязуется по заданию Заказчика оказать услуги: "
          "{{service_description}}, а Заказчик обязуется принять и оплатить "
          "оказанные услуги на условиях настоящего Договора.",
          size=10.5, space_after=6)

    _section_heading(doc, "2. СТОИМОСТЬ И ПОРЯДОК РАСЧЁТОВ")
    _para(doc, "2.1. Стоимость услуг по настоящему Договору составляет "
                "{{total}} {{currency}}.", size=10.5, space_after=4)
    _para(doc, "2.2. {{payment_terms}}", size=10.5, space_after=6)

    _section_heading(doc, "3. СРОК ДЕЙСТВИЯ ДОГОВОРА")
    _para(doc, "3.1. Настоящий Договор вступает в силу с {{start_date}} и "
                "действует до {{end_date}}.", size=10.5, space_after=6)

    for heading, text in _SERVICE_CONTRACT_BOILERPLATE:
        _section_heading(doc, heading)
        _para(doc, text, size=10.5, space_after=6)

    _section_heading(doc, "7. РЕКВИЗИТЫ И ПОДПИСИ СТОРОН")
    _parties_block(doc, left_label="ИСПОЛНИТЕЛЬ", right_label="ЗАКАЗЧИК",
                    right_extra=["Директор: {{client_director_name}}"])
    _signatures_block(doc, [
        ("Исполнитель:", "director_name"),
        ("Заказчик:", "client_director_name"),
    ])
    out = BytesIO(); doc.save(out); return out.getvalue()


def build_supply_contract() -> bytes:
    doc = Document()
    _set_margins(doc.sections[0])
    _title_block(doc, "ДОГОВОР ПОСТАВКИ", "{{document_number}}", "{{document_date}}")
    _para(doc, "г. {{city}}", size=10.5, space_after=10)
    _para(doc,
          "{{company_name}}, именуемое в дальнейшем «Поставщик», в лице "
          "{{director_name}}, действующего на основании Устава, с одной "
          "стороны, и {{client_name}}, именуемое в дальнейшем «Покупатель», "
          "в лице {{client_director_name}}, действующего на основании "
          "Устава, с другой стороны, совместно именуемые «Стороны», "
          "заключили настоящий Договор о нижеследующем:",
          size=10.5, space_after=10)

    _section_heading(doc, "1. ПРЕДМЕТ ДОГОВОРА")
    _para(doc, "1.1. Поставщик обязуется передать в собственность Покупателя "
                "товар согласно спецификации ниже, а Покупатель обязуется "
                "принять и оплатить товар на условиях настоящего Договора.",
          size=10.5, space_after=6)
    _table(doc, loop_var="items", columns=[
        ("№",                "{{ loop.index }}",       1.2),
        ("Наименование товара", "{{ it.name }}",       7.3),
        ("Ед.",              "шт",                     1.5),
        ("Кол-во",           "{{ it.qty }}",            1.5),
        ("Цена",             "{{ it.price_fmt }}",      2.5),
        ("Сумма",            "{{ it.total_fmt }}",      3.0),
    ])
    _para(doc, "", space_after=6)
    _totals_block(doc, with_vat=True)

    _section_heading(doc, "2. ЦЕНА И ПОРЯДОК РАСЧЁТОВ")
    _para(doc, "2.1. Общая стоимость товара по настоящему Договору составляет "
                "{{total}} {{currency}}, в т.ч. НДС {{vat_percent}}% — "
                "{{vat}} {{currency}}.", size=10.5, space_after=4)
    _para(doc, "2.2. {{payment_terms}}", size=10.5, space_after=6)

    _section_heading(doc, "3. УСЛОВИЯ ПОСТАВКИ")
    _para(doc, "3.1. {{delivery_terms}}", size=10.5, space_after=4)
    _para(doc, "3.2. Адрес поставки: {{delivery_address}}.", size=10.5, space_after=4)
    _para(doc, "3.3. Срок поставки: {{delivery_date}}.", size=10.5, space_after=6)

    for heading, text in _SUPPLY_CONTRACT_BOILERPLATE:
        _section_heading(doc, heading)
        _para(doc, text, size=10.5, space_after=6)

    _section_heading(doc, "7. РЕКВИЗИТЫ И ПОДПИСИ СТОРОН")
    _parties_block(doc, left_label="ПОСТАВЩИК", right_label="ПОКУПАТЕЛЬ",
                    right_extra=["Директор: {{client_director_name}}"])
    _signatures_block(doc, [
        ("Поставщик:", "director_name"),
        ("Покупатель:", "client_director_name"),
    ])
    _para(doc, "", space_after=10)
    _para(doc, "Банковские реквизиты Поставщика", size=10, bold=True)
    _para(doc, "{{company_name}} · БИН {{company_bin}}", size=10)
    _para(doc, "{{company_bank}}", size=10)
    out = BytesIO(); doc.save(out); return out.getvalue()


def build_reconciliation_act() -> bytes:
    doc = Document()
    _set_margins(doc.sections[0])
    _title_block(doc, "АКТ СВЕРКИ ВЗАИМОРАСЧЁТОВ", "{{document_number}}", "{{document_date}}",
                 subtitle="за период с {{period_start}} по {{period_end}}")
    _para(doc,
          "{{company_name}} (БИН {{company_bin}}) и {{client_name}} "
          "(БИН {{client_bin}}) составили настоящий акт о состоянии "
          "взаиморасчётов по данным бухгалтерского учёта Сторон:",
          size=10.5, space_after=10)
    _para(doc, "Входящее сальдо на {{period_start}}: {{opening_balance}} {{currency}}",
          size=10.5, bold=True, space_after=8)
    _table(doc, loop_var="operations", columns=[
        ("№",         "{{ loop.index }}",  1.2),
        ("Дата",      "{{ it.date }}",     2.3),
        ("Документ",  "{{ it.doc_ref }}",  5.5),
        ("Дебет",     "{{ it.debit_fmt }}", 3.0),
        ("Кредит",    "{{ it.credit_fmt }}", 3.0),
        ("Сальдо",    "{{ it.balance_fmt }}", 2.0),
    ])
    _para(doc, "", space_after=8)
    _para(doc, "Исходящее сальдо на {{period_end}}: {{closing_balance}} {{currency}}",
          size=10.5, bold=True, space_after=2)
    _para(doc, "Сумма прописью: {{closing_balance_words}}", size=10, italic=True, space_after=10)
    _para(doc, "Настоящий акт подтверждает отсутствие разногласий по "
                "указанным суммам между Сторонами на дату его подписания.",
          size=10, space_after=8)
    _parties_block(doc, left_label="ОРГАНИЗАЦИЯ 1", right_label="ОРГАНИЗАЦИЯ 2",
                    right_extra=["Руководитель: {{client_director_name}}"])
    _signatures_block(doc, [
        ("От организации 1:", "director_name"),
        ("От организации 2:", "client_director_name"),
    ])
    out = BytesIO(); doc.save(out); return out.getvalue()


def build_hr_order() -> bytes:
    doc = Document()
    _set_margins(doc.sections[0])
    _title_block(doc, "ПРИКАЗ", "{{order_number}}", "{{document_date}}",
                 subtitle="О приёме на работу")
    _para(doc, "г. {{city}}", size=10.5, space_after=12)
    _para(doc, "ПРИКАЗЫВАЮ:", size=11, bold=True, space_after=8)
    _para(doc, "1. Принять {{employee_name}} на должность «{{employee_position}}» "
                "в структурное подразделение «{{employee_department}}» с "
                "{{hire_date}} г.", size=10.5, space_after=6)
    _para(doc, "2. Установить оклад в размере {{salary}} тенге в месяц.",
          size=10.5, space_after=6)
    _para(doc, "3. Испытательный срок: {{probation_period}}.",
          size=10.5, space_after=6)
    _para(doc, "4. Приказ довести до сведения бухгалтерии для начисления "
                "заработной платы.", size=10.5, space_after=10)
    _para(doc, "Основание: трудовой договор от {{document_date}} г.",
          size=10, italic=True, space_after=16)
    _signatures_block(doc, [
        ("Руководитель:", "director_name"),
    ])
    _para(doc, "", space_after=10)
    _para(doc, "С приказом ознакомлен(а):", size=10, space_after=20)
    _para(doc, "_______________________ / {{employee_name}}      "
                "Дата: _______________", size=10)
    out = BytesIO(); doc.save(out); return out.getvalue()


_EMPLOYMENT_CONTRACT_BOILERPLATE = [
    ("5. ПРАВА И ОБЯЗАННОСТИ СТОРОН",
     "5.1. Работник обязуется добросовестно исполнять свои трудовые "
     "обязанности, соблюдать трудовую дисциплину и правила внутреннего "
     "трудового распорядка Работодателя. 5.2. Работодатель обязуется "
     "предоставить Работнику работу, обусловленную настоящим Договором, "
     "обеспечить условия труда и своевременно выплачивать заработную "
     "плату в полном размере."),
    ("6. ОТВЕТСТВЕННОСТЬ СТОРОН",
     "6.1. Стороны несут ответственность за неисполнение или ненадлежащее "
     "исполнение обязательств по настоящему Договору в соответствии с "
     "трудовым законодательством Республики Казахстан."),
    ("7. ЗАКЛЮЧИТЕЛЬНЫЕ ПОЛОЖЕНИЯ",
     "7.1. Изменение условий настоящего Договора допускается только по "
     "соглашению Сторон в письменной форме, за исключением случаев, "
     "предусмотренных трудовым законодательством Республики Казахстан. "
     "7.2. Настоящий Договор составлен в двух экземплярах, имеющих "
     "одинаковую юридическую силу, по одному для каждой из Сторон."),
]


def build_employment_contract() -> bytes:
    doc = Document()
    _set_margins(doc.sections[0])
    _title_block(doc, "ТРУДОВОЙ ДОГОВОР", "{{document_number}}", "{{document_date}}")
    _para(doc, "г. {{city}}", size=10.5, space_after=10)
    _para(doc,
          "{{company_name}}, именуемое в дальнейшем «Работодатель», в лице "
          "{{director_name}}, действующего на основании Устава, с одной "
          "стороны, и {{employee_name}}, именуемый(ая) в дальнейшем "
          "«Работник», с другой стороны, совместно именуемые «Стороны», "
          "заключили настоящий трудовой договор о нижеследующем:",
          size=10.5, space_after=10)

    _section_heading(doc, "1. ПРЕДМЕТ ДОГОВОРА")
    _para(doc, "1.1. Работник принимается на должность «{{employee_position}}» "
                "в структурное подразделение «{{employee_department}}».",
          size=10.5, space_after=4)
    _para(doc, "1.2. Настоящий Договор является {{contract_term}}.",
          size=10.5, space_after=4)
    _para(doc, "1.3. Дата начала работы: {{work_start_date}}.",
          size=10.5, space_after=6)

    _section_heading(doc, "2. ОПЛАТА ТРУДА")
    _para(doc, "2.1. Работнику устанавливается должностной оклад в размере "
                "{{salary}} тенге в месяц.", size=10.5, space_after=4)
    _para(doc, "2.2. {{pay_dates}}", size=10.5, space_after=6)

    _section_heading(doc, "3. РЕЖИМ РАБОЧЕГО ВРЕМЕНИ И ВРЕМЕНИ ОТДЫХА")
    _para(doc, "3.1. {{work_schedule}}", size=10.5, space_after=4)
    _para(doc, "3.2. Работнику предоставляется ежегодный оплачиваемый "
                "трудовой отпуск продолжительностью {{vacation_days}} "
                "календарных дней.", size=10.5, space_after=6)

    _section_heading(doc, "4. ИСПЫТАТЕЛЬНЫЙ СРОК")
    _para(doc, "4.1. {{probation_period}}", size=10.5, space_after=6)

    for heading, text in _EMPLOYMENT_CONTRACT_BOILERPLATE:
        _section_heading(doc, heading)
        _para(doc, text, size=10.5, space_after=6)

    _section_heading(doc, "8. РЕКВИЗИТЫ И ПОДПИСИ СТОРОН")
    _parties_block(doc, left_label="РАБОТОДАТЕЛЬ", right_label="РАБОТНИК",
                    right_base=["{{employee_name}}", "ИИН: {{employee_iin}}",
                                "{{employee_address}}"],
                    right_extra=["Документ: {{employee_id_doc}}"])
    _signatures_block(doc, [
        ("Работодатель:", "director_name"),
        ("Работник:", "employee_name"),
    ])
    out = BytesIO(); doc.save(out); return out.getvalue()


# ── Installer ──────────────────────────────────────────────────────────────

TEMPLATES = [
    ("trust_letter",         "Доверенность на получение ТМЦ",         build_trust_letter),
    ("contract_services",    "Договор оказания услуг",                build_service_contract),
    ("contract_supply",      "Договор поставки",                      build_supply_contract),
    ("act_reconciliation",   "Акт сверки взаиморасчётов",             build_reconciliation_act),
    ("hr_order",             "Приказ о приёме на работу",             build_hr_order),
    ("employment_contract",  "Трудовой договор",                      build_employment_contract),
]


def _build_mappings(kind: str) -> dict[str, dict]:
    return {
        key: {"source_label": key, "source_kind": "canonical",
              "confidence": 0.95, "confirmed": True}
        for key in CANONICAL_MAPPING_SETS[kind]
    }


async def install() -> None:
    storage = get_storage()
    async with SessionLocal() as session:
        company = (await session.scalars(select(Company).limit(1))).first()
        if not company:
            print("no company found — seed first"); return
        print(f"installing into company={company.id} ({company.name})")

        for kind, display_name, builder in TEMPLATES:
            await session.execute(
                delete(Template).where(
                    Template.company_id == company.id,
                    Template.name == display_name,
                ),
            )
            await session.flush()

            blob = builder()
            file_hash = hashlib.sha256(blob).hexdigest()
            tpl_id = f"tpl_{uuid.uuid4().hex[:10]}"
            storage_key = f"{company.id}/templates/{tpl_id}/{display_name}.docx"
            storage.put(storage_key, blob, content_type=(
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ))

            mappings = _build_mappings(kind)
            t = Template(
                id=tpl_id, company_id=company.id,
                name=display_name, kind=kind, format="docx",
                body=storage_key, is_default=True,
                status="VERIFIED", original_filename=f"{display_name}.docx",
                file_hash=file_hash, language="ru", version=1,
                detected_fields={
                    "placeholders": list(mappings.keys()),
                    "semantic": {"kind": kind, "commercial": True},
                    "built_by": "install_kz_legal_templates.py",
                },
                detected_tables=[],
                mappings=mappings,
                confidence=0.95,
                validation_score=95, validation_warnings=[],
                industry="general",
            )
            session.add(t)
            print(f"  [{kind:20s}] VERIFIED  {display_name}  "
                  f"({len(blob)} bytes, {len(mappings)} mappings)")

        await session.commit()
        print(f"\nDone. Installed {len(TEMPLATES)} KZ legal/HR templates.")


if __name__ == "__main__":
    asyncio.run(install())
