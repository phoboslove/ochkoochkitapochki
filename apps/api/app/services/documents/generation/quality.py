"""Render QA — Phase 1.

Inspects the *actual* rendered artifact (DOCX bytes / PDF bytes / HTML) plus
the canonical context that fed the renderer, and returns a structured
quality report.

Deductions are intentionally conservative — false-positives are worse than
false-negatives because they block real approvals. Each rule documents its
weight so the score is auditable.

Threshold convention (used by the approval gate):
  score >= 80   : OK            → approval PENDING (normal flow)
  60 <= score < 80 : WARNING    → approval PENDING + warning badge
  score < 60    : BLOCKED       → approval status BLOCKED; needs operator
"""
from __future__ import annotations

import re
from dataclasses import dataclass, field
from io import BytesIO
from typing import Any

from app.services.documents.generation.required_fields import (
    KINDS_WITH_TOTAL_ITEMS_CHECK, required_fields_for,
)

# A placeholder is anything still looking like {{ … }} or {% … %} in the output.
_LEFTOVER_PLACEHOLDER_RE = re.compile(r"\{\{[^}]{1,80}\}\}|\{%[^%]{1,80}%\}")

QUALITY_BLOCK_THRESHOLD = 60
QUALITY_WARN_THRESHOLD  = 80


@dataclass
class QualityIssue:
    code: str            # short machine code, e.g. "leftover_placeholder"
    message: str         # human description for the operator
    severity: str        # "error" (blocks) | "warning" | "info"
    weight: int          # points subtracted from 100
    where: str | None = None   # optional location: "page 2 / table 1 / row 3"


@dataclass
class QualityReport:
    score: int
    status: str           # "ok" | "warning" | "blocked"
    issues: list[QualityIssue] = field(default_factory=list)
    stats: dict[str, Any] = field(default_factory=dict)

    def as_dict(self) -> dict[str, Any]:
        return {
            "score": self.score,
            "status": self.status,
            "blocking": self.status == "blocked",
            "issue_count": len(self.issues),
            "issues": [
                {"code": i.code, "message": i.message, "severity": i.severity,
                 "weight": i.weight, "where": i.where}
                for i in self.issues
            ],
            "stats": self.stats,
        }


# ── Public ──────────────────────────────────────────────────────────────────

def check_render_quality(
    *,
    kind: str,
    canonical: dict[str, Any],
    rendered_bytes: bytes,
    rendered_ext: str,           # docx | xlsx | pdf | html
    pdf_bytes: bytes | None,
    template_format: str | None,
    pipeline_warnings: list[str] | None = None,
) -> QualityReport:
    issues: list[QualityIssue] = []
    stats: dict[str, Any] = {
        "rendered_bytes": len(rendered_bytes or b""),
        "rendered_ext":   rendered_ext,
        "pdf_bytes":      len(pdf_bytes or b""),
        "has_pdf":        bool(pdf_bytes),
    }

    # 1) Required canonical fields — declarative per kind (required_fields.py).
    # No if/elif here by design: adding a new kind means adding one entry to
    # KIND_REQUIRED_FIELDS, not touching this function.
    for rf in required_fields_for(kind):
        v = canonical.get(rf.key)
        is_missing = (not v) if rf.check == "list" else v in (None, "", "—", 0)
        if is_missing:
            issues.append(QualityIssue(
                code=f"missing_{rf.key}", severity=rf.severity, weight=rf.weight,
                message=f"{rf.label} is empty in the populated context.",
            ))

    # 2) Item table sanity — only for kinds that actually render a generic
    # {{items}} table with a single grand total (see KINDS_WITH_TOTAL_ITEMS_CHECK).
    # Kinds like act_reconciliation/hr_order have their own shape and are
    # covered entirely by their declarative required-field list above.
    if kind in KINDS_WITH_TOTAL_ITEMS_CHECK:
        total_v = canonical.get("total_raw")
        if total_v in (None, "", 0):
            issues.append(QualityIssue(
                code="missing_total_raw", severity="error", weight=20,
                message="Total amount is empty in the populated context.",
            ))
        items = canonical.get("items") or []
        stats["item_count"] = len(items)
        if not items:
            # Items missing isn't always fatal (некоторые акты — текстовые), но снимаем баллы.
            issues.append(QualityIssue(
                code="empty_items", severity="warning", weight=10,
                message="No line items resolved — table will render empty.",
            ))
        else:
            item_total = 0.0
            for idx, it in enumerate(items, start=1):
                if not (it.get("name") or "").strip():
                    issues.append(QualityIssue(
                        code="item_missing_name", severity="warning", weight=5,
                        where=f"row {idx}",
                        message=f"Row {idx} has no name — table cell will be blank.",
                    ))
                try:
                    item_total += float(it.get("total") or 0)
                except (TypeError, ValueError):
                    issues.append(QualityIssue(
                        code="item_bad_total", severity="warning", weight=5,
                        where=f"row {idx}",
                        message=f"Row {idx} total is not numeric: {it.get('total')!r}.",
                    ))
            stats["items_sum"] = round(item_total, 2)
            # Compare against canonical total when both present (within 1 KZT tolerance).
            try:
                grand = float(canonical.get("subtotal_raw") or canonical.get("total_raw") or 0)
                if item_total and grand and abs(item_total - grand) > 1.0:
                    issues.append(QualityIssue(
                        code="totals_mismatch", severity="error", weight=15,
                        message=(f"Sum of line items ({item_total:.2f}) does not match "
                                  f"document subtotal/total ({grand:.2f})."),
                    ))
            except (TypeError, ValueError):
                pass

    # 3) Render-artifact level checks.
    if rendered_ext == "docx":
        _check_docx(rendered_bytes, issues, stats)
        # Row-count consistency: the largest table is almost always the items
        # table; its data-row count should equal the number of canonical
        # items we expected to render. Surfaces collapse bugs (parser found
        # N items, render produced 1 generic row) and expansion bugs alike.
        # Only meaningful for kinds that have a generic items table at all.
        if kind in KINDS_WITH_TOTAL_ITEMS_CHECK:
            rendered_rows = stats.get("docx_items_table_data_rows")
            expected = stats.get("item_count") or 0
            if rendered_rows is not None and expected and rendered_rows != expected:
                issues.append(QualityIssue(
                    code="items_row_count_mismatch", severity="error", weight=20,
                    where=f"rendered={rendered_rows} expected={expected}",
                    message=(
                        f"Items table rendered {rendered_rows} data row(s) but the "
                        f"parsed intent expected {expected}. The proposal and the "
                        "rendered document show different data — investigate "
                        "tool/pipeline item override priority."
                    ),
                ))
    elif rendered_ext == "xlsx":
        _check_xlsx(rendered_bytes, issues, stats)
    elif rendered_ext == "html":
        _check_text(rendered_bytes.decode("utf-8", errors="ignore"), issues, stats)

    # 4) PDF artifact — size + page-count sanity. Required for trust gate.
    if pdf_bytes is None:
        # Templates with no available PDF converter are downgraded but not blocked.
        issues.append(QualityIssue(
            code="pdf_missing", severity="warning", weight=15,
            message="PDF export not produced — preview will be DOCX/HTML only.",
        ))
    else:
        _check_pdf(pdf_bytes, issues, stats)

    # 5) Pass-through pipeline warnings (e.g. "no VERIFIED template — used fallback").
    for w in pipeline_warnings or []:
        # The fallback warning is informational, not blocking.
        sev, weight = ("warning", 3) if "fallback" in w else ("warning", 5)
        issues.append(QualityIssue(
            code="pipeline_warning", severity=sev, weight=weight, message=w,
        ))

    # ── Score & status ────────────────────────────────────────────────────
    deductions = sum(i.weight for i in issues)
    score = max(0, 100 - deductions)
    if score < QUALITY_BLOCK_THRESHOLD:
        status = "blocked"
    elif score < QUALITY_WARN_THRESHOLD:
        status = "warning"
    else:
        status = "ok"
    return QualityReport(score=score, status=status, issues=issues, stats=stats)


# ── Per-format inspectors ───────────────────────────────────────────────────

def _check_docx(blob: bytes, issues: list[QualityIssue], stats: dict[str, Any]) -> None:
    """Walk the rendered DOCX, looking for leftover placeholders + empty cells."""
    try:
        from docx import Document as DocxDocument
    except Exception as exc:  # noqa: BLE001
        issues.append(QualityIssue(
            code="docx_inspector_unavailable", severity="warning", weight=2,
            message=f"DOCX QA skipped: python-docx not importable ({exc!s}).",
        ))
        return
    try:
        doc = DocxDocument(BytesIO(blob))
    except Exception as exc:  # noqa: BLE001
        issues.append(QualityIssue(
            code="docx_parse_failed", severity="error", weight=40,
            message=f"Rendered DOCX is not a valid Office Open XML file: {exc!s}",
        ))
        return

    leftover, empty_cells, total_cells, table_count, row_count = 0, 0, 0, 0, 0
    # Track the table that looks most like the items table (most rows or
    # contains a header cell with one of the item-table column labels)
    # so we can compare its data-row count to the expected item count.
    biggest_table_rows = 0
    items_table_data_rows: int | None = None
    _items_header_markers = (
        "наимен", "кол-во", "цена", "сумма",
        "item", "qty", "price", "amount",
    )
    for para in doc.paragraphs:
        if _LEFTOVER_PLACEHOLDER_RE.search(para.text):
            leftover += 1
    for t in doc.tables:
        table_count += 1
        rows_in_table = 0
        looks_like_items = False
        for row in t.rows:
            row_count += 1
            rows_in_table += 1
            row_has_value = False
            for cell in row.cells:
                total_cells += 1
                txt = cell.text or ""
                if rows_in_table == 1:
                    low = txt.lower()
                    if any(m in low for m in _items_header_markers):
                        looks_like_items = True
                if _LEFTOVER_PLACEHOLDER_RE.search(txt):
                    leftover += 1
                if txt.strip():
                    row_has_value = True
                else:
                    empty_cells += 1
            if not row_has_value:
                # Fully empty rows in the middle of a table are usually leftover
                # repeating-row anchors — flag as a warning but only once per table.
                issues.append(QualityIssue(
                    code="empty_table_row", severity="warning", weight=2,
                    where=f"table {table_count}",
                    message=f"Table {table_count} contains an empty row "
                            "(repeating-row template not consumed?).",
                ))
        # Header-only table — data rows = total rows minus the header.
        # We only commit a row-count signal when the table clearly LOOKS
        # like an items table (header markers matched). A "biggest table
        # by default" fallback was tried earlier but produced false
        # positives on inventory-style legacy templates where the largest
        # table is the inventory list, not the line-item summary.
        if looks_like_items:
            items_table_data_rows = max(0, rows_in_table - 1)
        if rows_in_table > biggest_table_rows:
            biggest_table_rows = rows_in_table
    stats.update({
        "docx_tables": table_count, "docx_rows": row_count,
        "docx_cells": total_cells, "docx_empty_cells": empty_cells,
        "docx_leftover_placeholders": leftover,
        "docx_items_table_data_rows": items_table_data_rows,
    })
    if leftover:
        issues.append(QualityIssue(
            code="leftover_placeholder", severity="error", weight=min(40, 8 + leftover * 4),
            message=f"{leftover} unresolved placeholder(s) still visible in the DOCX "
                    "(e.g. {{client_name}} not mapped).",
        ))
    # Very small DOCX = almost certainly empty body.
    if len(blob) < 2_000:
        issues.append(QualityIssue(
            code="docx_too_small", severity="warning", weight=10,
            message=f"DOCX is suspiciously small ({len(blob)} bytes).",
        ))


def _check_xlsx(blob: bytes, issues: list[QualityIssue], stats: dict[str, Any]) -> None:
    try:
        from openpyxl import load_workbook
    except Exception:  # noqa: BLE001
        return
    try:
        wb = load_workbook(BytesIO(blob), data_only=False)
    except Exception as exc:  # noqa: BLE001
        issues.append(QualityIssue(
            code="xlsx_parse_failed", severity="error", weight=40,
            message=f"Rendered XLSX is not a valid workbook: {exc!s}",
        ))
        return
    leftover = 0
    cells = 0
    for ws in wb.worksheets:
        for row in ws.iter_rows(values_only=True):
            for v in row:
                cells += 1
                if isinstance(v, str) and _LEFTOVER_PLACEHOLDER_RE.search(v):
                    leftover += 1
    stats.update({"xlsx_cells": cells, "xlsx_leftover_placeholders": leftover})
    if leftover:
        issues.append(QualityIssue(
            code="leftover_placeholder", severity="error", weight=min(40, 8 + leftover * 4),
            message=f"{leftover} unresolved placeholder(s) still visible in the XLSX.",
        ))


def _check_text(text: str, issues: list[QualityIssue], stats: dict[str, Any]) -> None:
    leftover = len(_LEFTOVER_PLACEHOLDER_RE.findall(text))
    stats["html_leftover_placeholders"] = leftover
    if leftover:
        issues.append(QualityIssue(
            code="leftover_placeholder", severity="error", weight=min(30, 6 + leftover * 3),
            message=f"{leftover} unresolved placeholder(s) in the HTML render.",
        ))


def _check_pdf(blob: bytes, issues: list[QualityIssue], stats: dict[str, Any]) -> None:
    # Header check.
    if not blob.startswith(b"%PDF"):
        issues.append(QualityIssue(
            code="pdf_malformed", severity="error", weight=40,
            message="PDF output does not start with %PDF — file is corrupt.",
        ))
        return
    if len(blob) < 1_000:
        issues.append(QualityIssue(
            code="pdf_too_small", severity="error", weight=25,
            message=f"PDF is only {len(blob)} bytes — almost certainly empty.",
        ))
    # Quick page-count heuristic without a heavyweight dependency.
    # `/Type /Page` (not /Pages) appears once per page in linearised PDFs.
    try:
        pages = blob.count(b"/Type /Page") - blob.count(b"/Type /Pages")
        pages = max(pages, 1)
        stats["pdf_pages"] = pages
        if pages == 0:
            issues.append(QualityIssue(
                code="pdf_zero_pages", severity="error", weight=30,
                message="PDF appears to contain zero pages.",
            ))
        if pages > 20:
            issues.append(QualityIssue(
                code="pdf_pagination_suspicious", severity="warning", weight=5,
                message=f"PDF rendered {pages} pages — unusually long for a single document.",
            ))
    except Exception:  # noqa: BLE001
        pass
