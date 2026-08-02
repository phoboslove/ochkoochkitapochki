"""Legacy KZ template adaptation layer — Phase 1.

Most uploaded Kazakhstan templates were authored by humans for humans:
free-text labels ("Поставщик:", "Заказчик:", "Итого:") followed by empty
slots, hand-drawn item tables, no Jinja markers anywhere. This module
**transiently** injects canonical ``{{placeholder}}`` anchors into a copy of
the template bytes so the existing renderer can fill them in.

KEY INVARIANTS:
  * The original ``Template.body`` storage object is NEVER mutated. The
    adapter receives raw bytes, returns new bytes for the renderer to consume.
  * Original formatting (fonts, merged cells, borders, alignment, signatures,
    images) is preserved — adaptation only touches the *text* of cells / runs
    that look like empty insertion slots.
  * If a label is ambiguous (multiple matches, value already populated, or
    the slot looks non-empty), we DO NOT inject — a warning diagnostic is
    recorded instead.
  * If the template already contains ``{{placeholders}}``, adaptation is a
    no-op (the template author already mapped fields).

Supported formats: docx (including docx-converted RTF), xlsx.
"""
from __future__ import annotations

import re
from io import BytesIO

from app.services.documents.generation.diagnostics import (
    STAGE_ADAPTATION, DiagnosticsCollector, RenderDiagnostic,
)
from app.services.templates.label_matcher import match_labels_batch

# Label → canonical-field matching (previously a hardcoded ANCHOR_LIBRARY
# substring table here) now goes through label_matcher.match_labels_batch —
# semantic embeddings first, falling back to placeholders.SYNONYMS substring
# matching when no AI key is configured. See label_matcher.py.

# Items table headers — each column maps to its canonical item field.
ITEMS_COLUMN_LIBRARY: dict[str, list[str]] = {
    "name":  ["наименование", "товар", "услуга", "услуги", "работа", "работы", "описание", "предмет"],
    "qty":   ["кол-во", "количество", "ед.изм", "ед. изм", "шт"],
    "price": ["цена", "стоимость единицы", "цена за единицу", "тариф"],
    "total": ["сумма", "стоимость", "итого", "итог"],
}

# Pattern detecting an "empty slot" — something the human left blank for ink.
EMPTY_SLOT_RE = re.compile(r"^[\s_\-\.…:№\(\)]*$")
# Pattern detecting an already-present Jinja placeholder.
HAS_PLACEHOLDER_RE = re.compile(r"\{\{[^}]+\}\}|\{%[^%]+%\}")
# Label-with-trailing-empty pattern: "Поставщик:" or "Поставщик ________"
LABEL_TAIL_RE = re.compile(r"^(?P<label>.+?)\s*[:\-]?\s*(?P<tail>[_\-\.…\s]{0,80})$")


async def adapt_template(
    *, content: bytes, format: str, existing_mappings: dict | None = None,
) -> tuple[bytes, DiagnosticsCollector]:
    """Inject virtual anchors in-memory. Returns (adapted_bytes, diagnostics).

    If the template is already fully mapped (``existing_mappings`` covers the
    canonical fields), or the format doesn't support adaptation, the original
    bytes are returned unchanged and the collector stays empty.
    """
    diag = DiagnosticsCollector()
    if not content:
        return content, diag
    fmt = (format or "").lower()
    if fmt in ("docx", "doc"):
        try:
            return await _adapt_docx(content, diag, existing_mappings or {})
        except Exception as exc:  # noqa: BLE001
            diag.warn(STAGE_ADAPTATION, "docx_adapter_failed",
                       f"DOCX adaptation skipped due to internal error: {exc!s}",
                       exception=str(exc)[:200], render_engine="python-docx",
                       suggested_fix="Re-upload the template and re-run analysis; "
                                     "if the issue persists, open the file in Word and resave.")
            return content, diag
    if fmt == "xlsx":
        try:
            return await _adapt_xlsx(content, diag, existing_mappings or {})
        except Exception as exc:  # noqa: BLE001
            diag.warn(STAGE_ADAPTATION, "xlsx_adapter_failed",
                       f"XLSX adaptation skipped due to internal error: {exc!s}",
                       exception=str(exc)[:200], render_engine="openpyxl")
            return content, diag
    # html / pdf / unsupported: no adaptation.
    return content, diag


# ── DOCX adapter ────────────────────────────────────────────────────────────

async def _adapt_docx(
    content: bytes, diag: DiagnosticsCollector, existing_mappings: dict,
) -> tuple[bytes, DiagnosticsCollector]:
    from docx import Document as DocxDocument

    doc = DocxDocument(BytesIO(content))

    # Skip outright if the template already has placeholders — author already
    # did the mapping work; nothing to fix.
    if _docx_already_templated(doc):
        diag.info(STAGE_ADAPTATION, "already_templated",
                   "Template already contains Jinja placeholders — adaptation skipped.")
        return content, diag

    canonicals_already_mapped = {
        k for k, v in existing_mappings.items() if v.get("confirmed") and v.get("source_label")
    }
    injected: set[str] = set(canonicals_already_mapped)

    # Collect every distinct candidate label up front so the whole document
    # resolves through ONE batched embeddings call instead of one round-trip
    # per paragraph/cell.
    labels: list[str] = []
    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if text and not HAS_PLACEHOLDER_RE.search(text):
            labels.append(text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = (cell.text or "").strip()
                if text and not HAS_PLACEHOLDER_RE.search(text):
                    labels.append(text)
    index = await match_labels_batch(labels)

    # 1) Paragraphs — single-line labels like "Поставщик: __________"
    for p_idx, para in enumerate(doc.paragraphs):
        _try_inject_paragraph(para, p_idx, injected, diag, index)

    # 2) Tables — labels in cells + items-table detection.
    for t_idx, table in enumerate(doc.tables):
        # 2a) Detect items header row.
        header_row, columns = _detect_items_header(table)
        if header_row is not None and header_row + 1 < len(table.rows):
            _inject_items_row(table, header_row, columns, t_idx, diag)
            injected.add("items")

        # 2b) Per-cell label adaptation.
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                _try_inject_cell(cell, row, c_idx, t_idx, r_idx, injected, diag, index)

    if not diag.entries:
        diag.info(STAGE_ADAPTATION, "no_anchors_matched",
                   "No semantic anchors matched — template will render as-is. "
                   "Renderer may fall back to HTML if required fields are missing.",
                   suggested_fix="Confirm template field mappings manually in the "
                                 "template detail page if rendering looks incomplete.")

    out = BytesIO()
    doc.save(out)
    return out.getvalue(), diag


def _docx_already_templated(doc) -> bool:
    """True iff any paragraph or cell already contains a Jinja placeholder."""
    for p in doc.paragraphs:
        if HAS_PLACEHOLDER_RE.search(p.text):
            return True
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                if HAS_PLACEHOLDER_RE.search(cell.text or ""):
                    return True
    return False


def _try_inject_paragraph(
    para, p_idx: int, injected: set[str], diag: DiagnosticsCollector,
    index: dict[str, tuple[str, float] | None],
) -> None:
    """If the paragraph looks like 'Label: ___', append the matching anchor."""
    text = (para.text or "").strip()
    if not text or HAS_PLACEHOLDER_RE.search(text):
        return
    match = index.get(text)
    if not match:
        return
    canonical, confidence = match
    label = text
    if canonical in injected:
        return
    # Only inject when the line is a label-with-empty-slot (e.g. ends in
    # underscores / colon / nothing). If the line already contains a value
    # ("Заказчик: TOO ABC"), abstain — we'd overwrite real data.
    m = LABEL_TAIL_RE.match(text)
    if m is None:
        return
    tail = m.group("tail") or ""
    label_part = m.group("label") or ""
    is_empty_slot = EMPTY_SLOT_RE.match(tail.strip("_-.…: ") or "") is not None
    has_value_after_label = bool(text[len(label_part):].strip(" :-_.…")) and not is_empty_slot
    if has_value_after_label:
        diag.warn(STAGE_ADAPTATION, "anchor_ambiguous",
                   f"Found '{label}' in paragraph {p_idx} but the slot already "
                   f"contains text — not auto-filling to avoid overwriting.",
                   failed_field=canonical, render_engine="python-docx",
                   extra={"paragraph_text": text[:120]},
                   suggested_fix="Confirm the mapping manually in template detail.")
        return
    # Append the canonical anchor as a new run so we don't disturb run styles
    # of the existing label text.
    placeholder = "{{" + canonical + "}}"
    sep = " " if text and not text.endswith(":") else " "
    para.add_run(sep + placeholder)
    injected.add(canonical)
    diag.info(STAGE_ADAPTATION, "anchor_injected",
               f"Injected {{{{{canonical}}}}} after label '{label}' (paragraph {p_idx}).",
               failed_field=canonical, render_engine="python-docx",
               extra={"label": label, "confidence": confidence})


def _try_inject_cell(
    cell, row, c_idx: int, t_idx: int, r_idx: int,
    injected: set[str], diag: DiagnosticsCollector,
    index: dict[str, tuple[str, float] | None],
) -> None:
    """If a cell contains only a label and its right-neighbour is empty,
    place the placeholder in the neighbour."""
    text = (cell.text or "").strip()
    if not text or HAS_PLACEHOLDER_RE.search(text):
        return
    match = index.get(text)
    if not match:
        return
    canonical, confidence = match
    label = text
    if canonical in injected:
        return
    # Only treat as a label cell when the cell text is short — accounting
    # labels ("Поставщик", "БИН заказчика", "Итого к оплате") are all well
    # under this, so a longer cell almost certainly already has a value
    # concatenated in (embeddings match on meaning, not position, so they
    # can't tell "just a label" from "label: value" the way a substring
    # match against a short keyword used to).
    if len(text) > 32:
        # Looks like the cell contains a value too — skip.
        return
    if c_idx + 1 >= len(row.cells):
        diag.warn(STAGE_ADAPTATION, "anchor_no_value_cell",
                   f"Label '{label}' found in last column of table {t_idx} row {r_idx} — "
                   "no adjacent cell to place the value.",
                   failed_field=canonical, table_index=t_idx, row_index=r_idx,
                   render_engine="python-docx",
                   suggested_fix="Re-design the template so the label is followed by a value cell.")
        return
    neighbour = row.cells[c_idx + 1]
    neighbour_text = (neighbour.text or "").strip()
    if neighbour_text and not EMPTY_SLOT_RE.match(neighbour_text):
        diag.warn(STAGE_ADAPTATION, "anchor_ambiguous",
                   f"Label '{label}' (table {t_idx} row {r_idx}) has a neighbour cell "
                   "with existing content — not overwriting.",
                   failed_field=canonical, table_index=t_idx, row_index=r_idx,
                   extra={"neighbour_text": neighbour_text[:80]},
                   suggested_fix="Confirm the mapping manually if this is the correct slot.")
        return
    placeholder = "{{" + canonical + "}}"
    # Write into the first paragraph of the neighbour cell, preserving cell formatting.
    target_para = neighbour.paragraphs[0] if neighbour.paragraphs else neighbour.add_paragraph()
    # Clear any underscore-y blank text in the slot before injecting.
    for r in list(target_para.runs):
        r.text = ""
    target_para.add_run(placeholder)
    injected.add(canonical)
    diag.info(STAGE_ADAPTATION, "anchor_injected",
               f"Injected {{{{{canonical}}}}} into table {t_idx} row {r_idx} after label '{label}'.",
               failed_field=canonical, table_index=t_idx, row_index=r_idx,
               render_engine="python-docx",
               extra={"label": label, "confidence": confidence})


# ── DOCX items-table detection ──────────────────────────────────────────────

def _detect_items_header(table) -> tuple[int | None, dict[int, str]]:
    """Return ``(header_row_idx, {column_index: canonical_item_field})``.

    A row counts as an items-table header when at least 2 distinct canonical
    item fields are recognised across its cells. Caller fills the row below
    (header_row + 1) with placeholders + a docxtpl ``{%tr for it in items %}``
    directive.
    """
    for r_idx, row in enumerate(table.rows):
        columns: dict[int, str] = {}
        for c_idx, cell in enumerate(row.cells):
            txt = (cell.text or "").strip().lower()
            if not txt:
                continue
            for field, keywords in ITEMS_COLUMN_LIBRARY.items():
                if any(k in txt for k in keywords):
                    columns[c_idx] = field
                    break
        if len(set(columns.values())) >= 2:
            return r_idx, columns
    return None, {}


def _inject_items_row(
    table, header_row: int, columns: dict[int, str], t_idx: int, diag: DiagnosticsCollector,
) -> None:
    """Fill row header_row+1 with item placeholders and wrap it in a docxtpl
    ``{%tr for ... %} ... {%tr endfor %}`` directive so the renderer expands
    one row per item while preserving the original row's borders/styles."""
    items_row = table.rows[header_row + 1]
    cells = list(items_row.cells)
    if not cells:
        return
    # Clear existing placeholder-y content first.
    for cell in cells:
        for p in cell.paragraphs:
            for r in list(p.runs):
                r.text = ""

    # Map column → field-expression (use formatted variants for price/total).
    field_expr = {
        "name":  "{{ it.name }}",
        "qty":   "{{ it.qty }}",
        "price": "{{ it.price_fmt }}",
        "total": "{{ it.total_fmt }}",
    }
    for c_idx, cell in enumerate(cells):
        field = columns.get(c_idx)
        if not field:
            continue
        target = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        target.add_run(field_expr[field])

    # Wrap the row with tr directives (docxtpl looks for `{%tr ... %}` anywhere
    # in the row's text). Put the for in the first cell and endfor in the last.
    first_para = cells[0].paragraphs[0] if cells[0].paragraphs else cells[0].add_paragraph()
    first_para.runs[0].text = "{%tr for it in items %}" + first_para.runs[0].text \
        if first_para.runs else first_para.add_run("{%tr for it in items %}").text
    if first_para.runs and not first_para.runs[0].text.startswith("{%tr"):
        # Prepend directive run.
        first_para.add_run("{%tr for it in items %}")
    last_para = cells[-1].paragraphs[-1] if cells[-1].paragraphs else cells[-1].add_paragraph()
    last_para.add_run("{%tr endfor %}")

    diag.info(STAGE_ADAPTATION, "items_table_detected",
               f"Items table detected in table {t_idx}: header at row {header_row}, "
               f"columns mapped: {list(columns.values())}. Row {header_row+1} "
               "will repeat per line item.",
               table_index=t_idx, render_engine="python-docx",
               extra={"columns": columns})


# ── XLSX adapter ────────────────────────────────────────────────────────────

async def _adapt_xlsx(
    content: bytes, diag: DiagnosticsCollector, existing_mappings: dict,
) -> tuple[bytes, DiagnosticsCollector]:
    from openpyxl import load_workbook

    wb = load_workbook(BytesIO(content))
    if _xlsx_already_templated(wb):
        diag.info(STAGE_ADAPTATION, "already_templated",
                   "Workbook already contains Jinja placeholders — adaptation skipped.")
        return content, diag

    injected: set[str] = {k for k, v in (existing_mappings or {}).items()
                            if v.get("confirmed") and v.get("source_label")}

    # Collect every distinct candidate label across all sheets up front so
    # the whole workbook resolves through ONE batched embeddings call.
    labels: list[str] = []
    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                if not isinstance(cell.value, str):
                    continue
                txt = cell.value.strip()
                if txt and not HAS_PLACEHOLDER_RE.search(txt):
                    labels.append(txt)
    index = await match_labels_batch(labels)

    for ws in wb.worksheets:
        # 1) Label cells.
        for row in ws.iter_rows():
            for cell in row:
                if not isinstance(cell.value, str):
                    continue
                txt = cell.value.strip()
                if not txt or HAS_PLACEHOLDER_RE.search(txt):
                    continue
                match = index.get(txt)
                if not match or match[0] in injected:
                    continue
                canonical, confidence = match
                label = txt
                # Labels are short ("БИН поставщика", "Итого к оплате"); a
                # longer cell almost certainly already has a value in it —
                # embeddings match on meaning, not position, so they can't
                # tell "just a label" from "label: value" by themselves.
                if len(txt) > 32:
                    continue
                # Target: cell to the right (most common KZ layout).
                target = ws.cell(row=cell.row, column=cell.column + 1)
                target_val = (target.value or "").strip() if isinstance(target.value, str) else target.value
                if target_val not in (None, "") and not (
                    isinstance(target_val, str) and EMPTY_SLOT_RE.match(target_val)
                ):
                    diag.warn(STAGE_ADAPTATION, "anchor_ambiguous",
                               f"Label '{label}' at {ws.title}!{cell.coordinate}: neighbour "
                               f"{target.coordinate} already has value — not overwriting.",
                               failed_field=canonical, render_engine="openpyxl",
                               extra={"sheet": ws.title, "cell": cell.coordinate})
                    continue
                target.value = "{{" + canonical + "}}"
                injected.add(canonical)
                diag.info(STAGE_ADAPTATION, "anchor_injected",
                           f"Injected {{{{{canonical}}}}} at {ws.title}!{target.coordinate} "
                           f"(label '{label}' at {cell.coordinate}).",
                           failed_field=canonical, render_engine="openpyxl",
                           extra={"sheet": ws.title, "label": label,
                                  "label_cell": cell.coordinate,
                                  "value_cell": target.coordinate, "confidence": confidence})

        # 2) Items header detection — first row with ≥ 2 column matches.
        items_header = _detect_xlsx_items_header(ws)
        if items_header is not None:
            header_row, columns = items_header
            _inject_xlsx_items_row(ws, header_row, columns, diag)
            injected.add("items")

    out = BytesIO()
    wb.save(out)
    return out.getvalue(), diag


def _xlsx_already_templated(wb) -> bool:
    for ws in wb.worksheets:
        for row in ws.iter_rows(values_only=True):
            for v in row:
                if isinstance(v, str) and HAS_PLACEHOLDER_RE.search(v):
                    return True
    return False


def _detect_xlsx_items_header(ws) -> tuple[int, dict[int, str]] | None:
    for row in ws.iter_rows():
        columns: dict[int, str] = {}
        for cell in row:
            if not isinstance(cell.value, str):
                continue
            txt = cell.value.strip().lower()
            for field, keywords in ITEMS_COLUMN_LIBRARY.items():
                if any(k in txt for k in keywords):
                    columns[cell.column] = field
                    break
        if len(set(columns.values())) >= 2:
            return row[0].row, columns
    return None


def _inject_xlsx_items_row(
    ws, header_row: int, columns: dict[int, str], diag: DiagnosticsCollector,
) -> None:
    target_row = header_row + 1
    field_expr = {
        "name":  "{{items.0.name}}",
        "qty":   "{{items.0.qty}}",
        "price": "{{items.0.price_fmt}}",
        "total": "{{items.0.total_fmt}}",
    }
    for col, field in columns.items():
        ws.cell(row=target_row, column=col).value = field_expr[field]
    diag.info(STAGE_ADAPTATION, "items_table_detected",
               f"Items header detected on sheet {ws.title!r} at row {header_row}; "
               f"row {target_row} will expand per item.",
               render_engine="openpyxl",
               extra={"sheet": ws.title, "header_row": header_row, "columns": columns})


