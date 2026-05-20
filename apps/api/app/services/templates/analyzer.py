"""Deterministic template analyzer.

Extracts:
  * paragraphs / text labels
  * table structures with column headers
  * Jinja-style {{placeholders}} already present
  * detected language (RU/KZ/EN heuristic)
  * document kind classification (invoice/act/nakladnaya/...)

NO LLM calls. NO autonomous rewriting. The output feeds the suggestion engine
and the mapping UI, where humans confirm before anything is rendered.
"""
from __future__ import annotations

import hashlib
import re
from dataclasses import asdict, dataclass, field
from io import BytesIO
from typing import Any

from app.services.templates.placeholders import CANONICAL, suggest


_PLACEHOLDER_RE = re.compile(r"\{\{\s*([\w\.\-]+)(?:\s*\|[^}]+)?\s*\}\}")

KIND_HINTS: list[tuple[str, str]] = [
    # Invoices
    ("счет на оплату",                          "invoice"),
    ("счёт на оплату",                          "invoice"),
    ("счёт-фактура",                            "invoice"),
    ("счет-фактура",                            "invoice"),
    ("шот-фактура",                             "invoice"),     # KZ
    ("invoice",                                 "invoice"),
    # Acts
    ("акт выполненных",                         "act"),
    ("акт сдачи",                               "act"),
    ("акт оказанных",                           "act"),
    ("акт о выявленных",                        "act"),
    ("акт приемки",                             "act"),
    ("акт приёмки",                             "act"),
    ("акт инвентаризации",                      "act"),
    ("акт списания",                            "act"),
    ("акт сверки",                              "act"),
    ("акт о ",                                  "act"),
    ("қабылдау актісі",                         "act"),         # KZ
    ("түгендеу актісі",                         "act"),         # KZ
    ("act of",                                  "act"),
    # Nakladnaya / waybill
    ("товарно-транспортная",                    "nakladnaya"),
    ("товарная накладная",                      "nakladnaya"),
    ("накладная на",                            "nakladnaya"),
    ("накладная",                               "nakladnaya"),
    ("жүкқұжат",                                "nakladnaya"),
    ("nakladnaya",                              "nakladnaya"),
    # Contract
    ("трудовой договор",                        "contract"),
    ("договор",                                 "contract"),
    ("контракт",                                "contract"),
    ("соглашение",                              "contract"),
    ("шарт",                                    "contract"),
    ("contract",                                "contract"),
    ("agreement",                               "contract"),
    # Proposal
    ("коммерческое предложение",                "proposal"),
    ("ценовое предложение",                     "proposal"),
    ("commercial offer",                        "proposal"),
    ("proposal",                                "proposal"),
    # Accounting forms — fold into act (most are signed statements/lists)
    ("доверенность",                            "act"),
    ("сенімхат",                                "act"),
    ("инвентаризационная опись",                "act"),
    ("инвентарная карточка",                    "act"),
    ("инвентарный список",                      "act"),
]


@dataclass
class DetectedTable:
    sheet: str | None = None
    headers: list[str] = field(default_factory=list)
    rows_preview: list[list[str]] = field(default_factory=list)  # max 3 rows
    bbox: tuple[int, int, int, int] | None = None
    is_repeating: bool = True


@dataclass
class Analysis:
    file_hash: str
    format: str
    language: str
    kind: str
    paragraphs: list[str]
    placeholders: list[str]
    tables: list[DetectedTable]
    suggestions: dict[str, dict[str, Any]]
    confidence: float       # 0..1 average mapping confidence
    notes: list[str] = field(default_factory=list)

    def to_dict(self) -> dict:
        return {
            "file_hash": self.file_hash, "format": self.format,
            "language": self.language, "kind": self.kind,
            "paragraphs": self.paragraphs[:60],
            "placeholders": self.placeholders,
            "tables": [asdict(t) for t in self.tables],
            "suggestions": self.suggestions,
            "confidence": round(self.confidence, 3),
            "notes": self.notes,
        }


def analyze(content: bytes, *, mime: str, filename: str) -> Analysis:
    fmt = _format_of(mime, filename)
    file_hash = hashlib.sha256(content).hexdigest()
    notes: list[str] = []

    if fmt == "docx":
        paragraphs, tables = _analyze_docx(content, notes)
    elif fmt == "xlsx":
        paragraphs, tables = _analyze_xlsx(content, notes)
    elif fmt == "pdf":
        paragraphs, tables = _analyze_pdf(content, notes)
    elif fmt == "rtf":
        paragraphs, tables = _analyze_rtf(content, notes)
    elif fmt in ("doc", "xls"):
        notes.append(
            f"Legacy {fmt.upper()} (binary Office 97-2003) — convert to "
            f"{'DOCX' if fmt == 'doc' else 'XLSX'} for full analysis."
        )
        paragraphs, tables = [], []
    else:  # html / unknown
        text = content.decode("utf-8", errors="ignore")
        paragraphs = [p.strip() for p in re.split(r"\n+", text) if p.strip()][:200]
        tables = []

    placeholders = sorted({m.group(1) for p in paragraphs for m in _PLACEHOLDER_RE.finditer(p)})
    language = _detect_language(paragraphs)
    kind = _classify_kind(paragraphs)

    # Build suggested mappings from paragraphs + table headers.
    sources: list[tuple[str, str]] = []  # (label, source_kind)
    for p in paragraphs:
        sources.append((p, "paragraph"))
    for t in tables:
        for h in t.headers:
            sources.append((h, f"table_header:{t.sheet or 'doc'}"))

    suggestions: dict[str, dict[str, Any]] = {}
    for label, source_kind in sources:
        match = suggest(label)
        if not match: continue
        canonical, conf = match
        prev = suggestions.get(canonical)
        if prev and prev["confidence"] >= conf:
            continue
        suggestions[canonical] = {
            "source_label": _truncate(label),
            "source_kind":  source_kind,
            "confidence":   round(conf, 3),
            "confirmed":    False,
        }

    avg_conf = (sum(s["confidence"] for s in suggestions.values()) / max(1, len(suggestions))) if suggestions else 0.0

    if not placeholders and not tables and not suggestions:
        notes.append("No placeholders, table headers, or recognizable accounting terms found. "
                     "Manual mapping will be required.")

    return Analysis(
        file_hash=file_hash, format=fmt, language=language, kind=kind,
        paragraphs=paragraphs, placeholders=placeholders, tables=tables,
        suggestions=suggestions, confidence=avg_conf, notes=notes,
    )


# ── Format-specific extractors ──────────────────────────────────────────────

def _analyze_docx(content: bytes, notes: list[str]) -> tuple[list[str], list[DetectedTable]]:
    try:
        from docx import Document  # python-docx
    except ImportError:
        notes.append("python-docx not installed — DOCX analyzed as plain text.")
        text = content.decode("utf-8", errors="ignore")
        return ([t.strip() for t in re.split(r"\n+", text) if t.strip()], [])
    try:
        doc = Document(BytesIO(content))
    except Exception as exc:  # noqa: BLE001
        notes.append(f"DOCX parse failed: {exc}")
        return ([], [])

    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    tables: list[DetectedTable] = []
    for tbl in doc.tables:
        if not tbl.rows: continue
        headers = [c.text.strip() for c in tbl.rows[0].cells]
        rows_preview: list[list[str]] = []
        for r in tbl.rows[1:4]:
            rows_preview.append([c.text.strip() for c in r.cells])
        # Heuristic: repeating row table if >= 1 of these terms in headers.
        keywords = ("наимен", "услуг", "товар", "qty", "кол", "сумм", "цена", "price")
        is_rep = any(any(k in h.lower() for k in keywords) for h in headers)
        tables.append(DetectedTable(
            sheet=None, headers=headers, rows_preview=rows_preview, is_repeating=is_rep,
        ))
    return paragraphs, tables


def _analyze_xlsx(content: bytes, notes: list[str]) -> tuple[list[str], list[DetectedTable]]:
    try:
        from openpyxl import load_workbook
    except ImportError:
        notes.append("openpyxl not installed — XLSX cannot be deeply analyzed.")
        return ([], [])
    try:
        wb = load_workbook(BytesIO(content), data_only=False)
    except Exception as exc:  # noqa: BLE001
        notes.append(f"XLSX parse failed: {exc}")
        return ([], [])

    paragraphs: list[str] = []
    tables: list[DetectedTable] = []
    for ws in wb.worksheets:
        # Collect single-cell text labels (paragraphs).
        for row in ws.iter_rows(max_row=80, values_only=True):
            for cell in row:
                if isinstance(cell, str) and cell.strip():
                    paragraphs.append(cell.strip())

        # Detect tables: find row with 3+ short string headers, then peek down.
        for r_idx, row in enumerate(ws.iter_rows(max_row=80, values_only=True), start=1):
            headers = [str(v).strip() if v is not None else "" for v in row]
            non_empty = [h for h in headers if h]
            if len(non_empty) < 3: continue
            if not all(0 < len(h) < 40 for h in non_empty): continue
            # Find first non-empty contiguous block.
            if not any(any(k in h.lower() for k in ("наимен", "услуг", "товар", "qty", "кол", "цена", "price", "сумм")) for h in non_empty):
                continue
            preview: list[list[str]] = []
            for i, prow in enumerate(ws.iter_rows(min_row=r_idx + 1, max_row=r_idx + 3, values_only=True)):
                preview.append([str(v) if v is not None else "" for v in prow])
            tables.append(DetectedTable(sheet=ws.title, headers=headers, rows_preview=preview, is_repeating=True))
            break  # one table per sheet for now
    return paragraphs[:200], tables


def _analyze_pdf(content: bytes, notes: list[str]) -> tuple[list[str], list[DetectedTable]]:
    # Lightweight: extract text via pypdf when present; tables left for future OCR pipeline.
    try:
        from pypdf import PdfReader
    except ImportError:
        notes.append("pypdf not installed — PDF analyzed as a placeholder. Wire OCR provider for real PDFs.")
        return ([], [])
    try:
        reader = PdfReader(BytesIO(content))
        text = "\n".join((p.extract_text() or "") for p in reader.pages[:5])
        return ([t.strip() for t in re.split(r"\n+", text) if t.strip()][:200], [])
    except Exception as exc:  # noqa: BLE001
        notes.append(f"PDF parse failed: {exc}")
        return ([], [])


# ── Utilities ────────────────────────────────────────────────────────────────

def _format_of(mime: str, filename: str) -> str:
    name = (filename or "").lower()
    if name.endswith(".docx") or "wordprocessingml" in mime: return "docx"
    if name.endswith(".xlsx") or "spreadsheetml" in mime:    return "xlsx"
    if name.endswith(".pdf") or mime == "application/pdf":   return "pdf"
    if name.endswith(".rtf") or "rtf" in mime:               return "rtf"
    if name.endswith(".doc"):                                return "doc"
    if name.endswith(".xls"):                                return "xls"
    if name.endswith(".html") or mime == "text/html":        return "html"
    return "html"


def _analyze_rtf(content: bytes, notes: list[str]) -> tuple[list[str], list[DetectedTable]]:
    """RTF text extraction — line by line. Tables not reliably extractable."""
    try:
        from striprtf.striprtf import rtf_to_text
    except ImportError:
        notes.append("striprtf not installed — RTF body kept raw.")
        return ([], [])
    try:
        text = rtf_to_text(content.decode("utf-8", errors="ignore"), errors="ignore")
    except Exception as exc:  # noqa: BLE001
        notes.append(f"RTF parse failed: {exc}")
        return ([], [])
    paragraphs = [p.strip() for p in re.split(r"\n+", text) if p.strip()][:240]
    notes.append("RTF table structure is not preserved — saved as text-only template.")
    return paragraphs, []


def _detect_language(paragraphs: list[str]) -> str:
    sample = " ".join(paragraphs[:30]).lower()
    cyr = len(re.findall(r"[а-яё]", sample))
    lat = len(re.findall(r"[a-z]", sample))
    kz  = len(re.findall(r"[әіңғүұқөһ]", sample))
    if kz > 3 and cyr > lat: return "kk"
    if cyr > lat:            return "ru"
    return "en"


def _classify_kind(paragraphs: list[str]) -> str:
    blob = " ".join(paragraphs[:40]).lower()
    for token, kind in KIND_HINTS:
        if token in blob:
            return kind
    return "unknown"


def _truncate(s: str, n: int = 80) -> str:
    s = s.strip()
    return s if len(s) <= n else s[: n - 1] + "…"
